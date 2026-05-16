import streamlit as st
import pandas as pd
from docx import Document
import os

st.set_page_config(page_title="inspeccionesPC", layout="wide")

st.title("📋 InspeccionesPC")

# --- ARCHIVO DE DATOS ---
archivo = "datos.xlsx"

if os.path.exists(archivo):
    df = pd.read_excel(archivo)
else:
    df = pd.DataFrame(columns=["Cliente", "Dirección", "Estado"])

# ---------------------------
# FORMULARIO
# ---------------------------
st.subheader("➕ Nuevo registro")

col1, col2, col3 = st.columns(3)

with col1:
    nombre = st.text_input("Cliente")

with col2:
    direccion = st.text_input("Dirección")

with col3:
    estado = st.selectbox("Estado", ["Activo", "Pendiente"])

# BOTÓN GUARDAR
if st.button("💾 Guardar"):
    nuevo = pd.DataFrame([[nombre, direccion, estado]],
                         columns=["Cliente", "Dirección", "Estado"])
    
    df = pd.concat([df, nuevo], ignore_index=True)
    df.to_excel(archivo, index=False)

    st.success("Registro guardado ✅")

# ---------------------------
# TABLA
# ---------------------------
st.subheader("📊 Listado de inspecciones")

# filtro
busqueda = st.text_input("🔍 Buscar cliente")

if busqueda:
    filtro_df = df[df["Cliente"].str.contains(busqueda, case=False, na=False)]
else:
    filtro_df = df

st.dataframe(filtro_df, use_container_width=True)

# ---------------------------
# GENERAR WORD
# ---------------------------
st.subheader("📄 Generar documento")

cliente_seleccionado = st.selectbox("Selecciona cliente", df["Cliente"] if not df.empty else [])

if st.button("📄 Generar Word"):
    if cliente_seleccionado:
        fila = df[df["Cliente"] == cliente_seleccionado].iloc[0]

        doc = Document()
        doc.add_heading('Informe de inspección', 0)

        doc.add_paragraph(f'Cliente: {fila["Cliente"]}')
        doc.add_paragraph(f'Dirección: {fila["Dirección"]}')
        doc.add_paragraph(f'Estado: {fila["Estado"]}')

        doc.save("informe.docx")

        with open("informe.docx", "rb") as f:
            st.download_button("⬇ Descargar Word", f, file_name="informe.docx")
